from __future__ import annotations

import io
import json
import re
import time
from datetime import date, datetime
from typing import Optional

from langchain.chat_models import init_chat_model
from pydantic import BaseModel, Field

MAX_INVOICE_TEXT_CHARS = 12000


class ParsedInvoice(BaseModel):
    invoice_id: Optional[str] = Field(default=None, description="Invoice number or ID")
    client_name: Optional[str] = Field(default=None, description="Bill-to client name")
    client_email: Optional[str] = Field(default=None, description="Bill-to client email")
    invoice_amount: Optional[float] = Field(default=None, description="Total amount due or balance due")
    due_date: Optional[str] = Field(default=None, description="Due date in YYYY-MM-DD format")
    days_overdue: Optional[int] = Field(default=None, description="Days overdue as of today")
    jurisdiction: Optional[str] = Field(default=None, description="Client jurisdiction/state/country if present")
    confidence: float = Field(default=0.0, description="Extraction confidence from 0 to 1")
    warnings: list[str] = Field(default_factory=list)


def validate_invoice_upload(filename: str, content_type: str | None) -> str:
    name = filename.lower()
    content_type = content_type or ""
    if name.endswith(".pdf") or content_type == "application/pdf":
        return "pdf"
    if name.endswith(".docx") or content_type in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }:
        return "docx"
    if name.endswith(".doc"):
        raise ValueError("Legacy .doc files are not supported yet. Please upload a PDF or DOCX invoice.")
    raise ValueError("Unsupported invoice file. Please upload a PDF or DOCX invoice.")


def extract_text_from_invoice(content: bytes, file_type: str) -> str:
    if file_type == "pdf":
        return _extract_pdf_text(content)
    if file_type == "docx":
        return _extract_docx_text(content)
    raise ValueError("Unsupported invoice file type.")


def parse_invoice_text(text: str) -> ParsedInvoice:
    bounded = text[:MAX_INVOICE_TEXT_CHARS]
    fallback = _heuristic_parse(bounded)
    try:
        parsed = _llm_parse(bounded)
    except Exception as exc:
        fallback.warnings.append(_friendly_llm_warning(exc))
        parsed = fallback

    merged = _merge_with_fallback(parsed, fallback)
    merged.warnings.extend(_missing_field_warnings(merged))
    return merged


def parsed_invoice_payload(parsed: ParsedInvoice, raw_text: str) -> dict:
    missing = []
    for field in ("invoice_id", "client_name", "client_email", "invoice_amount", "days_overdue"):
        value = getattr(parsed, field)
        if value is None or value == "":
            missing.append(field)
    return {
        "extracted": parsed.model_dump(),
        "missing_fields": missing,
        "warnings": parsed.warnings,
        "text_length": min(len(raw_text), MAX_INVOICE_TEXT_CHARS),
    }


def _extract_pdf_text(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    parts = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text:
            parts.append(page_text)
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("No readable text was found in this PDF. Scanned/image-only PDFs are not supported yet.")
    return text[:MAX_INVOICE_TEXT_CHARS]


def _extract_docx_text(content: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(content))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("No readable text was found in this DOCX invoice.")
    return text[:MAX_INVOICE_TEXT_CHARS]


def _llm_parse(text: str) -> ParsedInvoice:
    model = init_chat_model("llama-3.3-70b-versatile", model_provider="groq", temperature=0)
    extractor = model.with_structured_output(ParsedInvoice)
    prompt = (
        "Extract invoice fields from the invoice text. Prefer bill-to/client details over sender details. "
        "For money, prefer labels like balance due, amount due, total due, or invoice total; do not use tax/subtotal alone. "
        "Return days_overdue based on due_date and today's date when possible. "
        f"Today's date is {date.today().isoformat()}.\n\n"
        f"Invoice text:\n{text}"
    )
    last_error: Exception | None = None
    for attempt in range(3):
        try:
            result = extractor.invoke(prompt)
            if isinstance(result, ParsedInvoice):
                return result
            if isinstance(result, dict):
                return ParsedInvoice(**result)
            return ParsedInvoice(**json.loads(str(result)))
        except Exception as exc:
            last_error = exc
            message = str(exc).lower()
            if "429" not in message and "rate_limit" not in message:
                break
            time.sleep(1.5 * (attempt + 1))
    raise last_error or RuntimeError("Invoice extraction failed")


def _heuristic_parse(text: str) -> ParsedInvoice:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    email_match = re.search(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}", text)
    invoice_id = _first_match(text, [
        r"invoice\s*(?:#|no\.?|number|id)?\s*[:\-]?\s*([A-Z0-9][A-Z0-9._/-]{1,})",
        r"\binv[-\s#:]?([A-Z0-9._/-]{2,})",
    ])
    amount = _extract_amount(text)
    due_date = _extract_due_date(text)
    client_name = _extract_client_name(lines)
    jurisdiction = _first_match(text, [
        r"(?:bill to|client|customer)[\s\S]{0,160}?\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b(?:,\s*(?:USA|US|United States))?",
    ])
    days_overdue = _days_overdue(due_date) if due_date else None
    return ParsedInvoice(
        invoice_id=invoice_id,
        client_name=client_name,
        client_email=email_match.group(0) if email_match else None,
        invoice_amount=amount,
        due_date=due_date,
        days_overdue=days_overdue,
        jurisdiction=jurisdiction,
        confidence=0.35,
        warnings=[],
    )


def _extract_amount(text: str) -> Optional[float]:
    labels = [
        "balance due",
        "amount due",
        "total due",
        "invoice total",
        "total amount",
        "grand total",
        "total",
    ]
    for label in labels:
        pattern = rf"{label}\s*[:\-]?\s*(?:USD|US\$|\$)?\s*([0-9][0-9,]*(?:\.[0-9]{{2}})?)"
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            return float(match.group(1).replace(",", ""))
    amounts = re.findall(r"(?:USD|US\$|\$)\s*([0-9][0-9,]*(?:\.[0-9]{2})?)", text)
    if amounts:
        values = [float(amount.replace(",", "")) for amount in amounts]
        return max(values)
    return None


def _extract_due_date(text: str) -> Optional[str]:
    match = re.search(r"due\s*date\s*[:\-]?\s*([A-Za-z]{3,9}\s+\d{1,2},?\s+\d{4}|\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}-\d{2}-\d{2})", text, re.I)
    if not match:
        return None
    raw = match.group(1).strip()
    for fmt in ("%Y-%m-%d", "%m/%d/%Y", "%d/%m/%Y", "%m-%d-%Y", "%d-%m-%Y", "%B %d, %Y", "%b %d, %Y", "%B %d %Y", "%b %d %Y"):
        try:
            return datetime.strptime(raw, fmt).date().isoformat()
        except ValueError:
            continue
    return None


def _extract_client_name(lines: list[str]) -> Optional[str]:
    for idx, line in enumerate(lines):
        if re.search(r"\b(bill to|client|customer)\b", line, re.I):
            same_line = re.sub(r".*?(bill to|client|customer)\s*[:\-]?", "", line, flags=re.I).strip()
            if same_line and "@" not in same_line:
                return same_line[:120]
            for candidate in lines[idx + 1:idx + 4]:
                if "@" not in candidate and not re.search(r"\$|invoice|date|due", candidate, re.I):
                    return candidate[:120]
    return None


def _first_match(text: str, patterns: list[str]) -> Optional[str]:
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            return match.group(1).strip()
    return None


def _days_overdue(due_date: str) -> Optional[int]:
    try:
        due = datetime.strptime(due_date, "%Y-%m-%d").date()
    except ValueError:
        return None
    return max(0, (date.today() - due).days)


def _merge_with_fallback(parsed: ParsedInvoice, fallback: ParsedInvoice) -> ParsedInvoice:
    data = parsed.model_dump()
    fallback_data = fallback.model_dump()
    for key, value in fallback_data.items():
        if key == "warnings":
            continue
        if data.get(key) in (None, "", 0.0) and value not in (None, ""):
            data[key] = value
    warnings = list(dict.fromkeys([*parsed.warnings, *fallback.warnings]))
    data["warnings"] = warnings
    if not data.get("days_overdue") and data.get("due_date"):
        data["days_overdue"] = _days_overdue(data["due_date"])
    return ParsedInvoice(**data)


def _missing_field_warnings(parsed: ParsedInvoice) -> list[str]:
    warnings = []
    labels = {
        "invoice_id": "Invoice ID was not found.",
        "client_name": "Client name was not found.",
        "client_email": "Client email was not found.",
        "invoice_amount": "Invoice amount was not found.",
        "days_overdue": "Due date/days overdue was not found.",
    }
    for field, message in labels.items():
        value = getattr(parsed, field)
        if value is None or value == "":
            warnings.append(message)
    return warnings


def _friendly_llm_warning(exc: Exception) -> str:
    message = str(exc).lower()
    if "429" in message or "rate_limit" in message:
        return "AI extraction is temporarily rate limited; basic document parsing was used instead."
    return "AI extraction failed; basic document parsing was used instead."
